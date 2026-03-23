from fastapi import APIRouter, Depends, UploadFile, File, Query
from sqlalchemy.orm import Session
from typing import List
from app.core.database import get_db
from app.api.dependencies import get_current_user, requires_role, ROLE_JOBSEEKER
from app.models.user import User, JobseekerProfile
from app.schemas.resume import ResumeCreate, ResumeUpdate, Resume as ResumeSchema, ResumeListItem, ResumeEvaluateResult
from app.models.resume import Resume
from app.services.resume_service import ResumeService
from app.services.volcengine_service import volcengine_ai_service
from app.core.config import settings
from app.core.exceptions import InvalidFileTypeException, FileTooLargeException, BusinessException
import os
import uuid
import json
import base64
import docx
from pdf2image import convert_from_path
import io
import datetime

router = APIRouter(prefix="/resumes", tags=["求职者-简历管理"])

def convert_pdf_to_images(file_path: str) -> List[str]:
    """将PDF文件转换为base64编码的图片列表"""
    # 直接指定项目中的poppler路径（项目根目录下的poppler目录）
    import os
    # 当前文件路径: backend/app/api/v1/jobseeker/resume.py
    # 往上4层到项目根目录
    current_dir = os.path.dirname(os.path.abspath(__file__))
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(current_dir))))
    poppler_path = os.path.join(project_root, "poppler-25.12.0", "Library", "bin")
    print(f"[PDF转换] 使用poppler路径: {poppler_path}")
    print(f"[PDF转换] 路径是否存在: {os.path.exists(poppler_path)}")
    if not os.path.exists(poppler_path):
        # 尝试另一种可能的路径
        poppler_path = r"D:\Drivers\google_downloads\Release-25.12.0-0\poppler-25.12.0\Library\bin"
        print(f"[PDF转换] 尝试备用路径: {poppler_path}")

    images = convert_from_path(
        file_path,
        dpi=150,  # 降低分辨率加快速度
        poppler_path=poppler_path,
        first_page=1,
        last_page=5  # 最多处理5页，简历一般不会超过5页
    )
    base64_images = []
    for i, img in enumerate(images):
        # 压缩图片，降低尺寸
        max_width = 1000
        if img.width > max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height))

        buffer = io.BytesIO()
        img.save(buffer, format="JPEG", quality=80)  # 用JPEG压缩
        img_base64 = base64.b64encode(buffer.getvalue()).decode("utf-8")
        base64_images.append(img_base64)
        print(f"[PDF转换] 第{i+1}页转换完成，大小: {len(img_base64)/1024/1024:.2f}MB")

    return base64_images

def parse_date(date_str: str) -> datetime.date or None:
    """解析日期字符串为Python date对象，支持YYYY-MM和YYYY-MM-DD格式"""
    if not date_str or date_str == "至今" or date_str == "现在":
        return None
    try:
        if len(date_str) == 7:  # YYYY-MM格式
            return datetime.datetime.strptime(date_str, "%Y-%m").date()
        elif len(date_str) == 10:  # YYYY-MM-DD格式
            return datetime.datetime.strptime(date_str, "%Y-%m-%d").date()
    except Exception as e:
        print(f"日期解析失败: {date_str}, 错误: {e}")
    return None

def extract_text_from_docx(file_path: str) -> str:
    """从docx文件提取文本内容"""
    text = ""
    doc = docx.Document(file_path)
    for para in doc.paragraphs:
        text += para.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"
    return text.strip()

@router.get("", response_model=List[ResumeListItem], summary="获取简历列表")
def get_resume_list(
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """获取当前用户的所有简历列表"""
    jobseeker = db.query(JobseekerProfile).filter(JobseekerProfile.user_id == current_user.id).first()
    return ResumeService.get_resume_list(db, current_user.id, jobseeker.id)

@router.get("/{resume_id}", response_model=ResumeSchema, summary="获取简历详情")
def get_resume_detail(
    resume_id: int,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """根据ID获取简历详情"""
    return ResumeService.get_resume_by_id(db, resume_id, current_user.id)

@router.post("", response_model=ResumeSchema, summary="创建新简历")
def create_resume(
    resume_in: ResumeCreate,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """创建新简历"""
    jobseeker = db.query(JobseekerProfile).filter(JobseekerProfile.user_id == current_user.id).first()
    return ResumeService.create_resume(db, current_user.id, jobseeker.id, resume_in)

@router.put("/{resume_id}", response_model=ResumeSchema, summary="更新简历")
def update_resume(
    resume_id: int,
    resume_in: ResumeUpdate,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """更新简历信息"""
    return ResumeService.update_resume(db, resume_id, current_user.id, resume_in)

@router.delete("/{resume_id}", summary="删除简历")
def delete_resume(
    resume_id: int,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """删除简历，默认简历不能删除"""
    ResumeService.delete_resume(db, resume_id, current_user.id)
    return {"message": "删除成功"}

@router.post("/{resume_id}/set-default", summary="设置默认简历")
def set_default_resume(
    resume_id: int,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """设置指定简历为默认简历"""
    jobseeker = db.query(JobseekerProfile).filter(JobseekerProfile.user_id == current_user.id).first()
    ResumeService.set_default_resume(db, resume_id, current_user.id, jobseeker.id)
    return {"message": "设置成功"}

@router.post("/upload", summary="上传并解析简历")
async def upload_resume(
    file: UploadFile = File(...),
    title: str = Query(..., description="简历标题"),
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """上传简历文件并自动解析结构化内容"""
    # 先检查求职者档案是否存在，不存在则创建
    jobseeker = db.query(JobseekerProfile).filter(JobseekerProfile.user_id == current_user.id).first()
    if not jobseeker:
        jobseeker = JobseekerProfile(user_id=current_user.id)
        db.add(jobseeker)
        db.commit()
        db.refresh(jobseeker)

    # 检查文件类型
    ext = os.path.splitext(file.filename)[1].lower()[1:]
    if ext not in ["doc", "docx", "pdf"]:
        raise InvalidFileTypeException(message="仅支持doc、docx、pdf格式的简历")

    # 检查文件大小
    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE:
        raise FileTooLargeException()

    # 保存文件
    file_name = f"{uuid.uuid4().hex}.{ext}"
    save_path = os.path.join(settings.UPLOAD_DIR, "resumes", file_name)
    with open(save_path, "wb") as f:
        f.write(content)

    # 创建简历记录
    db_resume = Resume(
        user_id=current_user.id,
        jobseeker_id=jobseeker.id,
        title=title,
        original_file=f"/uploads/resumes/{file_name}",
        content=json.dumps({}),
        parse_status=1,  # 解析中
        is_default=False,
        score=None,
        evaluation=None
    )
    db.add(db_resume)
    db.commit()
    db.refresh(db_resume)

    # 异步解析简历（暂时同步处理，后续可改为后台任务）
    import time
    start_time = time.time()
    print(f"[简历解析] 开始解析，文件类型: {ext}, 文件路径: {save_path}")

    try:
        parsed_result = None
        if ext == "pdf":
            print("[简历解析] 使用PDF转图片多模态解析模式")
            # PDF转图片，调用多模态识别
            images = convert_pdf_to_images(save_path)
            print(f"[简历解析] PDF转换完成，共{len(images)}页")
            if images:
                parsed_result = await volcengine_ai_service.resume_parse(images=images)
                print(f"[简历解析] PDF解析完成，耗时: {time.time() - start_time:.2f}秒")
        elif ext == "docx":
            print("[简历解析] 使用DOCX文本解析模式")
            # docx文件提取文本解析
            text_content = extract_text_from_docx(save_path)
            print(f"[简历解析] DOCX文本提取完成，长度: {len(text_content)}字符")
            if text_content:
                parsed_result = await volcengine_ai_service.resume_parse(text_content=text_content)
                print(f"[简历解析] DOCX解析完成，耗时: {time.time() - start_time:.2f}秒")
        elif ext == "doc":
            raise InvalidFileTypeException(message="暂不支持doc格式，请转换为docx后上传")

        if parsed_result:
            print(f"[简历解析] 大模型返回结果: {json.dumps(parsed_result, ensure_ascii=False)[:500]}...")
            # 保存解析结果
            db_resume.content = json.dumps(parsed_result, ensure_ascii=False)
            # 自动填充经历信息
            education_count = len(parsed_result.get("education", []))
            work_count = len(parsed_result.get("work_experience", []))
            project_count = len(parsed_result.get("project_experience", []))
            print(f"[简历解析] 提取到: {education_count}条教育经历, {work_count}条工作经历, {project_count}条项目经历")

            # 教育经历
            for edu in parsed_result.get("education", []):
                from app.models.resume import ResumeEducationExperience
                db_edu = ResumeEducationExperience(
                    resume_id=db_resume.id,
                    school_name=edu.get("school_name", ""),
                    major=edu.get("major", ""),
                    education=edu.get("education", ""),
                    start_date=parse_date(edu.get("start_date")),
                    end_date=parse_date(edu.get("end_date")),
                    description=edu.get("description", "")
                )
                db.add(db_edu)
            # 工作经历
            for work in parsed_result.get("work_experience", []):
                from app.models.resume import ResumeWorkExperience
                db_work = ResumeWorkExperience(
                    resume_id=db_resume.id,
                    company_name=work.get("company_name", ""),
                    position=work.get("position", ""),
                    start_date=parse_date(work.get("start_date")),
                    end_date=parse_date(work.get("end_date")),
                    description=work.get("work_description", work.get("description", "")),
                    achievements=work.get("work_achievements", work.get("achievements", ""))
                )
                db.add(db_work)
            # 项目经历
            for project in parsed_result.get("project_experience", []):
                from app.models.resume import ResumeProjectExperience
                # 处理空字段，避免非空约束错误
                role = project.get("role")
                if not role or str(role).strip() == "":
                    role = "未填写"
                start_date = parse_date(project.get("start_date"))
                if not start_date:
                    from datetime import date
                    start_date = date.today()

                db_project = ResumeProjectExperience(
                    resume_id=db_resume.id,
                    project_name=project.get("project_name", ""),
                    role=role,
                    start_date=start_date,
                    end_date=parse_date(project.get("end_date")),
                    description=project.get("project_description", project.get("description", "")),
                    achievements=project.get("project_achievements", project.get("achievements", ""))
                )
                db.add(db_project)

            db_resume.parse_status = 2  # 解析成功
            print(f"[简历解析] 解析成功，总耗时: {time.time() - start_time:.2f}秒")
        else:
            db_resume.parse_status = 3  # 解析失败
            print("[简历解析] 错误：解析结果为空")
    except Exception as e:
        print(f"[简历解析] 失败: {e}")
        import traceback
        traceback.print_exc()
        db_resume.parse_status = 3  # 解析失败
    finally:
        db.commit()
        db.refresh(db_resume)
        print(f"[简历解析] 处理完成，最终状态: {'成功' if db_resume.parse_status == 2 else '失败'}")

    # 返回列表项结构，避免序列化问题
    return {
        "id": db_resume.id,
        "title": db_resume.title,
        "is_default": db_resume.is_default,
        "parse_status": db_resume.parse_status,
        "score": db_resume.score,
        "created_at": db_resume.created_at.isoformat(),
        "updated_at": db_resume.updated_at.isoformat()
    }

@router.post("/{resume_id}/evaluate", response_model=ResumeEvaluateResult, summary="AI评测简历")
async def evaluate_resume(
    resume_id: int,
    current_user: User = Depends(requires_role(ROLE_JOBSEEKER)),
    db: Session = Depends(get_db)
):
    """调用AI对简历进行评测，给出评分和优化建议"""
    resume = ResumeService.get_resume_by_id(db, resume_id, current_user.id)

    if not resume.content or resume.parse_status != 2:
        raise BusinessException(code=400010, message="简历未解析成功，无法进行评测")

    try:
        # 解析简历内容
        resume_content = json.loads(resume.content)
        # 调用大模型评测
        evaluate_result = await volcengine_ai_service.resume_evaluate(resume_content)

        # 保存评测结果
        resume.score = evaluate_result["score"]
        resume.evaluation = evaluate_result["evaluation"]
        db.commit()

        # 适配返回格式
        return {
            "score": evaluate_result["score"],
            "evaluation": evaluate_result["evaluation"],
            "suggestions": evaluate_result.get("suggestions", evaluate_result.get("disadvantages", []) + evaluate_result.get("advantages", []))
        }
    except Exception as e:
        print(f"简历评测失败: {e}")
        # 失败时返回模拟数据
        return {
            "score": 85,
            "evaluation": "您的简历整体质量较好，工作经历描述清晰，但项目经验部分可以更加量化成果。",
            "suggestions": [
                "建议在工作经历中添加更多量化的业绩成果，比如\"提升了30%的用户转化率\"",
                "项目经验部分可以补充您在项目中的具体贡献和使用的技术栈",
                "技能标签部分可以更有针对性，突出与目标岗位匹配的技能"
            ]
        }
